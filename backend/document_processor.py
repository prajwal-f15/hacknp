"""
Document Processing Pipeline using LangGraph
Step-based execution with model switching and fallback logic
"""

from typing import TypedDict, Annotated, Literal, Callable, Dict, Any, Optional
from langgraph.graph import StateGraph, END
from langchain_core.messages import HumanMessage
import re
import os
from pathlib import Path
import logging

# File handling
from PIL import Image
import pytesseract
import PyPDF2
import docx

# ML/AI Libraries
import spacy
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ModelRouter:
    """Routes and switches between different AI models"""
    
    def __init__(self, models_dir: Path):
        self.models_dir = models_dir
        self.available_models = self._check_available_models()
        
    def _check_available_models(self) -> Dict[str, bool]:
        """Check which models are available"""
        models = {
            'tesseract': self._check_tesseract(),
            'spacy': self._check_spacy(),
            'regex': True,  # Always available
            'mistral': self._check_mistral()
        }
        return models
    
    def _check_tesseract(self) -> bool:
        try:
            pytesseract.get_tesseract_version()
            return True
        except:
            return False
    
    def _check_spacy(self) -> bool:
        try:
            nlp = spacy.load("en_core_web_sm")
            return True
        except:
            return False
    
    def _check_mistral(self) -> bool:
        mistral_path = self.models_dir / "Mistral-7B-Instruct"
        return mistral_path.exists()
    
    def route_to_ocr_model(self, state: Dict, progress_callback: Optional[Callable] = None) -> str:
        """Route to OCR model (Tesseract)"""
        if progress_callback:
            progress_callback("🔄 Switching to OCR Model (Tesseract)", 15)
        
        if self.available_models['tesseract']:
            logger.info("✓ Routing to Tesseract OCR model")
            return "tesseract"
        else:
            logger.warning("⚠ Tesseract not available, skipping OCR")
            return "skip"
    
    def route_to_ner_model(self, state: Dict, progress_callback: Optional[Callable] = None) -> str:
        """Route to Named Entity Recognition model (spaCy)"""
        if progress_callback:
            progress_callback("🔄 Switching to NER Model (spaCy)", 45)
        
        if self.available_models['spacy']:
            logger.info("✓ Routing to spaCy NER model")
            return "spacy"
        else:
            logger.warning("⚠ spaCy not available, skipping NER")
            return "skip"
    
    def route_to_regex_model(self, state: Dict, progress_callback: Optional[Callable] = None) -> str:
        """Route to Regex pattern extraction model"""
        if progress_callback:
            progress_callback("🔄 Switching to Regex Pattern Extractor", 60)
        
        logger.info("✓ Routing to Regex pattern extraction")
        return "regex"
    
    def route_to_llm_model(self, state: Dict, progress_callback: Optional[Callable] = None) -> str:
        """Route to Large Language Model (Mistral)"""
        if progress_callback:
            progress_callback("🔄 Switching to LLM Model (Mistral-7B)", 75)
        
        if self.available_models['mistral'] and state.get('use_ai_summary', False):
            logger.info("✓ Routing to Mistral-7B LLM")
            return "mistral"
        else:
            logger.info("→ Routing to Simple Summarization (fallback)")
            return "simple"
    
    def get_model_status(self) -> Dict[str, str]:
        """Get status of all models"""
        status = {}
        for model, available in self.available_models.items():
            status[model] = "✅ Ready" if available else "❌ Not Available"
        return status

# State definition for the processing pipeline
class DocumentState(TypedDict):
    """State for document processing workflow"""
    file_path: str
    file_type: str
    raw_text: str
    cleaned_text: str
    entities: dict
    structured_data: dict
    summary: str
    error: str
    processing_step: str
    use_ai_summary: bool


class DocumentProcessor:
    """Main document processor using LangGraph for orchestration"""
    
    def __init__(self, models_dir: Path):
        self.models_dir = models_dir
        self.graph = self._build_graph()
        
        # Model Router - Switches between different models
        self.model_router = ModelRouter(models_dir)
        
    def _build_graph(self) -> StateGraph:
        """Build the LangGraph workflow"""
        
        # Create the graph
        workflow = StateGraph(DocumentState)
        
        # Add nodes for each processing step
        workflow.add_node("extract_text", self._extract_text_node)
        workflow.add_node("clean_text", self._clean_text_node)
        workflow.add_node("analyze_entities", self._analyze_entities_node)
        workflow.add_node("extract_structured_data", self._extract_structured_data_node)
        workflow.add_node("summarize", self._summarize_node)
        
        # Define the flow
        workflow.set_entry_point("extract_text")
        
        # Linear flow with conditional routing
        workflow.add_edge("extract_text", "clean_text")
        workflow.add_edge("clean_text", "analyze_entities")
        workflow.add_edge("analyze_entities", "extract_structured_data")
        workflow.add_edge("extract_structured_data", "summarize")
        workflow.add_edge("summarize", END)
        
        return workflow.compile()
    
    # ============================================================================
    # NODE 1: Text Extraction (Microservice-like isolation)
    # ============================================================================
    def _extract_text_node(self, state: DocumentState) -> DocumentState:
        """Extract text from various document types - supports images, PDFs, DOCX, TXT, etc."""
        state["processing_step"] = "Extracting text..."
        
        try:
            file_type = state["file_type"].lower()
            
            # Image formats (PNG, JPG, JPEG, BMP, TIFF, GIF, WEBP) - Use OCR
            if file_type in ['png', 'jpg', 'jpeg', 'bmp', 'tiff', 'tif', 'gif', 'webp']:
                model = self.model_router.route_to_ocr_model(state)
                if model == "tesseract":
                    state["raw_text"] = self._extract_from_image(state["file_path"])
                    logger.info(f"✓ OCR extracted {len(state['raw_text'])} chars from image")
                else:
                    state["raw_text"] = "[OCR not available - Please install Tesseract]"
                    state["error"] = "OCR not configured"
                    
            # PDF documents
            elif file_type == 'pdf':
                state["raw_text"] = self._extract_from_pdf(state["file_path"])
                logger.info(f"✓ Extracted {len(state['raw_text'])} chars from PDF")
                
            # Word documents (DOCX, DOC)
            elif file_type in ['docx', 'doc']:
                state["raw_text"] = self._extract_from_docx(state["file_path"])
                logger.info(f"✓ Extracted {len(state['raw_text'])} chars from DOCX")
                
            # Text files (TXT, CSV, LOG, MD)
            elif file_type in ['txt', 'text', 'csv', 'log', 'md']:
                state["raw_text"] = self._extract_from_txt(state["file_path"])
                logger.info(f"✓ Extracted {len(state['raw_text'])} chars from text file")
                
            else:
                state["error"] = f"Unsupported file type: {file_type}. Supported: PNG, JPG, PDF, DOCX, TXT, CSV, BMP, TIFF, GIF"
                state["raw_text"] = ""
                logger.error(f"✗ Unsupported file type: {file_type}")
            
            # Ensure we always have text (even if empty)
            if "raw_text" not in state or state["raw_text"] is None:
                state["raw_text"] = ""
                
        except Exception as e:
            state["error"] = f"Text extraction error: {str(e)}"
            state["raw_text"] = ""
            logger.error(f"✗ Text extraction failed: {str(e)}")
            
        return state
    
    def _extract_from_image(self, file_path: str) -> str:
        """OCR microservice - supports PNG, JPG, JPEG, BMP, TIFF, GIF, WEBP"""
        try:
            import pytesseract
            from PIL import Image
            
            # Try to find Tesseract in common installation paths
            tesseract_paths = [
                r"D:\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                r"C:\Tesseract-OCR\tesseract.exe"
            ]
            
            tesseract_found = False
            for path in tesseract_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    tesseract_found = True
                    logger.info(f"✓ Tesseract found at {path}")
                    break
            
            if not tesseract_found:
                return "[OCR Error: Tesseract not installed. Please run models/tesseract-installer.exe]"
            
            # Open and process image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary (handles PNG with alpha, etc.)
            if image.mode not in ('RGB', 'L'):
                image = image.convert('RGB')
            
            # Perform OCR with English language
            text = pytesseract.image_to_string(image, lang='eng')
            
            if not text or len(text.strip()) == 0:
                return "[OCR Warning: No text detected in image]"
            
            return text
            
        except ImportError:
            return "[OCR Error: pytesseract or Pillow not installed. Run: pip install pytesseract pillow]"
        except Exception as e:
            return f"[OCR Error: {str(e)}]"
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """PDF extraction microservice - handles text PDFs and scanned PDFs"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                # If no text extracted (scanned PDF), suggest OCR
                if not text or len(text.strip()) < 50:
                    return "[PDF Warning: This appears to be a scanned PDF with no extractable text. Please save as image (PNG/JPG) and upload for OCR processing]"
                
                return text
                
        except ImportError:
            return "[PDF Error: PyPDF2 not installed. Run: pip install PyPDF2]"
        except Exception as e:
            return f"[PDF Error: {str(e)}. File may be corrupted or encrypted]"
    
    def _extract_from_docx(self, file_path: str) -> str:
        """DOCX extraction microservice - handles Word documents and tables"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            text = "\n".join(text_parts)
            
            if not text or len(text.strip()) == 0:
                return "[DOCX Warning: No text found in document]"
            
            return text
            
        except ImportError:
            return "[DOCX Error: python-docx not installed. Run: pip install python-docx]"
        except Exception as e:
            return f"[DOCX Error: {str(e)}. File may be corrupted or in unsupported format]"
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Text file extraction microservice - handles TXT, CSV, LOG, MD with various encodings"""
        # Try multiple encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    if text:
                        return text
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                return f"[Text Error: {str(e)}]"
        
        # If all encodings fail
        return "[Text Error: Unable to decode file. File may be binary or use unsupported encoding]"
    
    # ============================================================================
    # NODE 2: Text Cleaning (Preprocessing microservice)
    # ============================================================================
    def _clean_text_node(self, state: DocumentState) -> DocumentState:
        """Clean and preprocess text - remove templates, logos, PII, unnecessary info"""
        state["processing_step"] = "Cleaning text..."
        
        try:
            raw_text = state.get("raw_text", "")
            
            # ==================================================================
            # COMPREHENSIVE PII REMOVAL - Handles repeated label format
            # Format: "Label Label : : Value"
            # ==================================================================
            
            text = raw_text
            
            # 1. Names - handles "Beneficiary Name Beneficiary MALVI" and "Register Worker Name : : Self"
            text = re.sub(r'Beneficiary\s*Name\s*Beneficiary\s*[A-Z]+', '[NAME-REDACTED]', text)
            text = re.sub(r'(?i)(patient\s*name|beneficiary\s*name|register\s*worker\s*name)[^:]*:[^:]*:\s*[^\n]+', '[NAME-REDACTED]', text)
            text = re.sub(r'(?i)name\s*:\s*[A-Z]{3,}[A-Z\s]+', '[NAME-REDACTED]', text)
            
            # 2. Phone numbers - handles "Contact No Contact No : : 7057912840"
            text = re.sub(r'(?i)contact\s*no\s*contact\s*no[^:]*:[^:]*:\s*\d{10,15}', '[PHONE-REDACTED]', text)
            text = re.sub(r'(?i)(contact|phone|mobile)\s*(?:no\.?|number)?[^:]*:[^:]*:\s*\d{10,15}', '[PHONE-REDACTED]', text)
            text = re.sub(r'\b\d{10}\b', '[PHONE-REDACTED]', text)
            text = re.sub(r'[+]?\d{10,15}', '[PHONE-REDACTED]', text)
            
            # 3. Email addresses
            text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL-REDACTED]', text)
            
            # 4. Address - handles "Address Address : : House No.,P 131besa..."
            text = re.sub(r'(?i)address\s*address[^:]*:[^:]*:\s*House[^\n]+', '[ADDRESS-REDACTED]', text)
            text = re.sub(r'(?i)address[^:]*:[^:]*:\s*[^\n]+\d{6}', '[ADDRESS-REDACTED]', text)
            text = re.sub(r'House\s*No\.?[^,]+,[^,]+,\d{6}', '[ADDRESS-REDACTED]', text)
            
            # 5. Age - handles "Age (Yr) Age (Yr) : : 49" and "49Y/FEMALE"
            text = re.sub(r'(?i)age\s*\([^)]+\)\s*age\s*\([^)]+\)[^:]*:[^:]*:\s*\d+', '[AGE-REDACTED]', text)
            text = re.sub(r'(?i)age[^:]*:[^:]*:\s*\d+', '[AGE-REDACTED]', text)
            text = re.sub(r'\d{2}Y/(?:MALE|FEMALE)', '[AGE-REDACTED]', text)
            text = re.sub(r'Age/Gender\s*:[^:]*:\s*\d+Y/[A-Z]+', '[AGE-REDACTED]', text)
            
            # 6. Gender - handles "Gender Gender : : Female"
            text = re.sub(r'(?i)gender\s*gender[^:]*:[^:]*:\s*(?:Male|Female)', '[GENDER-REDACTED]', text)
            text = re.sub(r'(?i)gender[^:]*:[^:]*:\s*(?:Male|Female)', '[GENDER-REDACTED]', text)
            
            # 7. Registration/Patient IDs - handles "Registration Number Registration Number : : 313030009368"
            text = re.sub(r'(?i)registration\s*number\s*registration\s*number[^:]*:[^:]*:\s*[A-Z0-9]+', '[ID-REDACTED]', text)
            text = re.sub(r'(?i)(registration\s*number|patient\s*id)[^:]*:[^:]*:\s*[A-Z0-9]{8,}', '[ID-REDACTED]', text)
            text = re.sub(r'(?i)patient\s+id\s*:\s*CWH\d+', '[PATIENT-ID-REDACTED]', text)  # Specific format
            text = re.sub(r'\b[A-Z]{3}\d{11,15}\b', '[REPORT-ID-REDACTED]', text)
            text = re.sub(r'CWH\d+', '[REPORT-ID-REDACTED]', text)
            
            # 8. Location - handles "District District : : Nagpur", "Taluka Taluka : : Nagpur (urban)"
            text = re.sub(r'(?i)district\s*district[^:]*:[^:]*:\s*[^\n]+', '[LOCATION-REDACTED]', text)
            text = re.sub(r'(?i)taluka\s*taluka[^:]*:[^:]*:\s*[^\n]+', '[LOCATION-REDACTED]', text)
            text = re.sub(r'(?i)(district|taluka)[^:]*:[^:]*:\s*[^\n]+', '[LOCATION-REDACTED]', text)
            text = re.sub(r'Maharashtra\s+India', '[LOCATION-REDACTED]', text)  # State + Country
            text = re.sub(r'\bMaharashtra\b', '[LOCATION-REDACTED]', text)  # State name
            text = re.sub(r'\bIndia\b', '[LOCATION-REDACTED]', text)  # Country name
            
            # 9. Pincode - handles "Pincode Pincode : : 440037"
            text = re.sub(r'(?i)pincode\s*pincode[^:]*:[^:]*:\s*\d{6}', '[PINCODE-REDACTED]', text)
            text = re.sub(r'(?i)pincode[^:]*:[^:]*:\s*\d{6}', '[PINCODE-REDACTED]', text)
            text = re.sub(r'\b\d{6}\b', '[PINCODE-REDACTED]', text)
            
            # 10. Relation - handles "Relation With Registered Worker : : Self"
            text = re.sub(r'(?i)relation\s*with[^:]*:[^:]*:\s*[^\n]+', '[RELATION-REDACTED]', text)
            
            # 11. Doctor names - handles "Dr. Dr Abhishek Sundeepkumar Singh"
            text = re.sub(r'Dr\.\s*Dr\s+[A-Z][a-z]+\s+[A-Z][a-z]+\s+[A-Z][a-z]+', '[DOCTOR-REDACTED]', text)
            text = re.sub(r'Dr\.\s*[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?', '[DOCTOR-REDACTED]', text)
            text = re.sub(r'(?i)Dr\s+[A-Z][a-z]+\s+[A-Z][a-z]+', '[DOCTOR-REDACTED]', text)  # Without period
            text = re.sub(r'\b[A-Z][a-z]+\s+Nasre\b', '[DOCTOR-REDACTED]', text)  # Specific doctor names
            text = re.sub(r'\bNitesh\s+Nasre\b', '[DOCTOR-REDACTED]', text)
            text = re.sub(r'(?i)registration\s*no\s*:\s*\d{10}', '[DOCTOR-REG-REDACTED]', text)
            text = re.sub(r'MD\s+Pathology', '[TITLE-REDACTED]', text)  # Medical titles
            
            # 12. Dates - handles "Date Of Screening Date Of Screening : : 01/11/2025"
            text = re.sub(r'(?i)date\s*of\s*screening\s*date\s*of\s*screening[^:]*:[^:]*:\s*\d{2}/\d{2}/\d{4}', '[DATE-REDACTED]', text)
            text = re.sub(r'(?i)(registered\s*on|reported\s*on)[^:]*:[^:]*:\s*\d{2}/\d{2}/\d{4}\s*\d{2}:\d{2}\s*[ap]m', '[DATETIME-REDACTED]', text)
            text = re.sub(r'\d{2}/\d{2}/\d{4}', '[DATE-REDACTED]', text)
            
            # 13. Height/Weight - handles "Height (cm) Height (cm) : : 153"
            text = re.sub(r'(?i)(height|weight)\s*\([^)]+\)\s*\1\s*\([^)]+\)[^:]*:[^:]*:\s*\d+', '[PHYSICAL-DATA-REDACTED]', text)
            text = re.sub(r'(?i)(height|weight)[^:]*:[^:]*:\s*\d+', '[PHYSICAL-DATA-REDACTED]', text)
            
            # 14. Lab locations - "Processed At : PLOT NO A/45,Back side..."
            text = re.sub(r'(?i)processed\s*at\s*:[^\n]+MIDC[^\n]+', '[LAB-LOCATION-REDACTED]', text)
            text = re.sub(r'PLOT\s*NO\s*[^\n,]+,[^\n]+', '[LAB-LOCATION-REDACTED]', text)
            text = re.sub(r'(?i)processed\s*at\s*:[^\n]+', '[LAB-LOCATION-REDACTED]', text)
            
            # 15. Page numbers and camp references
            text = re.sub(r'Page\s*No\s*-\s*\d+', '', text)
            text = re.sub(r'D2D\s*Camp\s*/\s*\d+\s*/[^\n]+', '', text)
            
            # 16. Customer name - "Customer Name : MBOCWWB"
            text = re.sub(r'(?i)customer\s*name[^:]*:[^:]*:\s*[^\n]+', '[CUSTOMER-REDACTED]', text)
            
            logger.info(f"✓ PII removed from text ({len(raw_text)} → {len(text)} chars)")
            
            # Remove ONLY template placeholders (NOT our PII redaction markers)
            text = re.sub(r'\[Your Name\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[Your Company Name\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[Your Address\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[Your Email\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[Your Phone\]', '', text, flags=re.IGNORECASE)
            # REMOVED: text = re.sub(r'\[.*?\]', '', text) - This was removing our [REDACTED] markers!
            
            # Remove common template phrases
            text = re.sub(r'(?i)your\s+logo\s+here', '', text)
            text = re.sub(r'(?i)company\s+logo', '', text)
            text = re.sub(r'(?i)insert\s+logo', '', text)
            text = re.sub(r'(?i)logo\s+placeholder', '', text)
            
            # Remove form field labels that are just placeholders
            text = re.sub(r'(?i)enter\s+your\s+\w+', '', text)
            text = re.sub(r'(?i)type\s+here', '', text)
            text = re.sub(r'(?i)click\s+to\s+add', '', text)
            
            # Remove watermark-like text
            text = re.sub(r'(?i)draft|sample|template|specimen', '', text)
            
            # Remove extra whitespace
            text = re.sub(r'\s+', ' ', text)
            
            # Remove special characters but keep basic punctuation
            text = re.sub(r'[^\w\s.,!?;:()\-\'\"/\n]', '', text)
            
            # Remove multiple dots/dashes
            text = re.sub(r'\.{3,}', '...', text)
            text = re.sub(r'\-{2,}', '-', text)
            
            # Remove lines with only special characters
            lines = text.split('\n')
            cleaned_lines = [
                line.strip() 
                for line in lines 
                if line.strip() and not all(c in '.-_=*#@' for c in line.strip())
            ]
            text = '\n'.join(cleaned_lines)
            
            # Remove very short lines (likely noise)
            lines = text.split('\n')
            cleaned_lines = [
                line for line in lines 
                if len(line.strip()) > 3 or line.strip() in ['.', '!', '?']
            ]
            text = ' '.join(cleaned_lines)
            
            # Remove repeated header/footer artifacts (same text appearing 3+ times)
            words = text.split()
            if len(words) > 20:
                # Check for repeated sequences
                from collections import Counter
                bigrams = [' '.join(words[i:i+5]) for i in range(len(words)-4)]
                repeated = [phrase for phrase, count in Counter(bigrams).items() if count >= 3]
                for phrase in repeated:
                    text = text.replace(phrase, '', 1)  # Keep one occurrence only
            
            # Final cleanup
            text = text.strip()
            text = re.sub(r'\s+', ' ', text)
            
            state["cleaned_text"] = text
            logger.info(f"✓ Text cleaned: {len(raw_text)} → {len(text)} characters")
            
        except Exception as e:
            state["error"] = f"Text cleaning error: {str(e)}"
            state["cleaned_text"] = state.get("raw_text", "")
            logger.error(f"✗ Text cleaning error: {str(e)}")
            
        return state
    
    # ============================================================================
    # NODE 3: Entity Analysis (spaCy microservice)
    # ============================================================================
    def _analyze_entities_node(self, state: DocumentState) -> DocumentState:
        """Extract named entities using spaCy with model routing"""
        state["processing_step"] = "Analyzing entities..."
        
        try:
            # Route to spaCy NER model
            model = self.model_router.route_to_ner_model(state)
            
            if model == "spacy":
                import spacy
                nlp = spacy.load("en_core_web_sm")
                
                text = state.get("cleaned_text", "")
                doc = nlp(text)
                
                entities = {
                    'persons': [],
                    'locations': [],
                    'organizations': [],
                    'dates': [],
                    'money': []
                }
                
                for ent in doc.ents:
                    if ent.label_ == "PERSON":
                        entities['persons'].append(ent.text)
                    elif ent.label_ in ["GPE", "LOC"]:
                        entities['locations'].append(ent.text)
                    elif ent.label_ == "ORG":
                        entities['organizations'].append(ent.text)
                    elif ent.label_ == "DATE":
                        entities['dates'].append(ent.text)
                    elif ent.label_ == "MONEY":
                        entities['money'].append(ent.text)
            else:
                state["entities"] = {}
                logger.warning("⚠ spaCy not available, skipping NER")
            
        except Exception as e:
            state["error"] = f"Entity analysis error: {str(e)}"
            state["entities"] = {}
            logger.error(f"✗ Entity analysis error: {str(e)}")
            
        return state
    
    # ============================================================================
    # NODE 4: Structured Data Extraction (Regex microservice)
    # ============================================================================
    def _extract_structured_data_node(self, state: DocumentState) -> DocumentState:
        """Extract structured data using regex patterns with model routing"""
        state["processing_step"] = "Extracting structured data..."
        
        try:
            # Route to Regex pattern model
            model = self.model_router.route_to_regex_model(state)
            
            import sys
            sys.path.insert(0, str(self.models_dir))
            from regex_patterns import IndianDataExtractor
            
            text = state.get("cleaned_text", "")
            extractor = IndianDataExtractor()
            data = extractor.extract_all(text)
            
            state["structured_data"] = data
            logger.info(f"✓ Extracted structured data: {len(data.get('phone_numbers', []))} phones, {len(data.get('aadhaar_numbers', []))} Aadhaar numbers")
            
        except Exception as e:
            state["error"] = f"Structured data extraction error: {str(e)}"
            state["structured_data"] = {}
            logger.error(f"✗ Structured data extraction error: {str(e)}")
            
        return state
    
    # ============================================================================
    # NODE 5: Summarization with Router (Model switching)
    # ============================================================================
    def _summarize_node(self, state: DocumentState) -> DocumentState:
        """Summarize using LLM model router with fallback"""
        state["processing_step"] = "Generating summary..."
        
        try:
            text = state.get("cleaned_text", "")
            use_ai = state.get("use_ai_summary", True)
            
            # Route to LLM model
            model = self.model_router.route_to_llm_model(state)
            
            if model == "mistral":
                summary = self._try_ai_summary(text)
                if summary:  # If we got a valid summary (not None)
                    state["summary"] = f"""📋 AI-Generated Summary (Mistral-7B)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{summary}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ Generated using Mistral-7B-Instruct (100% offline)"""
                    logger.info("✓ Generated AI summary with Mistral-7B")
                else:
                    # Fallback to simple summarization
                    state["summary"] = f"""📝 Simple Summary (Extractive)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{self._simple_summary(text)}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ℹ️ Mistral-7B unavailable, using extractive summary"""
                    logger.info("→ Fallback to simple summary")
            else:
                state["summary"] = f"""📝 Simple Summary (Extractive)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{self._simple_summary(text)}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ℹ️ AI summary not requested"""
                logger.info("✓ Generated simple summary")
                
        except Exception as e:
            logger.error(f"Summarization node error: {str(e)}")
            # Always provide a summary, never crash
            state["error"] = f"Summarization error: {str(e)}"
            state["summary"] = f"""📝 Simple Summary (Fallback)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{self._simple_summary(state.get("cleaned_text", ""))}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ Error occurred, using basic summary"""
            
        return state
    
    def _clear_memory(self):
        """Clear GPU and CPU memory by unloading models and running garbage collection"""
        try:
            import gc
            import torch
            
            # Force garbage collection
            gc.collect()
            
            # Clear CUDA cache if available
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                torch.cuda.synchronize()
                logger.info("✓ Cleared GPU memory")
            
            logger.info("✓ Memory cleared")
        except Exception as e:
            logger.warning(f"⚠ Memory clear warning: {str(e)}")
    
    def _try_lm_studio_summary(self, text: str) -> str:
        """Try to use LM Studio local API (much faster and more stable)"""
        try:
            import requests
            import json
            
            # LM Studio API endpoint - user's custom address
            url = "http://192.168.56.1:12345/v1/chat/completions"
            
            logger.info("📡 Connecting to LM Studio at http://192.168.56.1:12345...")
            
            # Create medical report prompt with STRICT privacy rules
            prompt = f"""You are a medical report analyzer. Generate a SHORT medical summary.

⚠️ CRITICAL - NEVER INCLUDE:
❌ Names (patient, doctor, beneficiary, worker)
❌ Age, gender, physical measurements
❌ Phone numbers, contact info
❌ Addresses, locations, districts, pincodes
❌ Registration numbers, IDs, report codes
❌ Dates, times
❌ Lab names, hospital locations

✅ ONLY INCLUDE:
• Medical test results (normal/abnormal ranges)
• Important health findings
• Recommendations for treatment

📝 Write ONLY 1-2 SHORT paragraphs in simple language.

Medical Report Data:
{text[:2000]}

Generate medical summary focusing ONLY on test results and health recommendations. NO personal information."""
            
            payload = {
                "model": "mistral-7b-instruct",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 300,
                "stream": False
            }
            
            headers = {"Content-Type": "application/json"}
            
            # Send request to LM Studio
            response = requests.post(url, json=payload, headers=headers, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                summary = result["choices"][0]["message"]["content"].strip()
                logger.info("✓ Summary generated via LM Studio")
                return summary
            else:
                logger.warning(f"⚠ LM Studio returned status {response.status_code}")
                return None
                
        except requests.exceptions.ConnectionError:
            logger.warning("⚠ LM Studio not running or not accessible at http://192.168.56.1:12345")
            return None
        except Exception as e:
            logger.warning(f"⚠ LM Studio error: {str(e)}")
            return None
    
    def _try_ai_summary(self, text: str, max_length: int = 300) -> str:
        """Try to use AI model - LM Studio first, then Hugging Face Mistral as fallback"""
        
        # OPTION 1: Try LM Studio first (faster and more stable)
        logger.info("🚀 Attempting LM Studio API...")
        lm_studio_summary = self._try_lm_studio_summary(text)
        if lm_studio_summary:
            return lm_studio_summary
        
        # OPTION 2: Fallback to Hugging Face Mistral
        logger.info("⚠ LM Studio unavailable, trying Hugging Face Mistral...")
        
        model = None
        tokenizer = None
        
        try:
            import torch
            from transformers import AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig
            import gc
            
            logger.info("🧹 Clearing memory before loading Mistral-7B...")
            # STEP 1: Clear all memory before loading Mistral
            self._clear_memory()
            
            logger.info("📦 Loading Mistral-7B model... (This may take 2-3 minutes)")
            logger.info("⏳ Please wait, loading tokenizer...")
            model_name = "mistralai/Mistral-7B-Instruct-v0.2"
            
            # Configure quantization
            quantization_config = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_compute_dtype=torch.float16,
                bnb_4bit_use_double_quant=True,
                bnb_4bit_quant_type="nf4"
            )
            
            tokenizer = AutoTokenizer.from_pretrained(model_name)
            logger.info("✓ Tokenizer loaded")
            logger.info("⏳ Loading model weights with 4-bit quantization... (1-2 minutes)")
            
            model = AutoModelForCausalLM.from_pretrained(
                model_name,
                quantization_config=quantization_config,
                device_map="auto",
                dtype=torch.float16,
                low_cpu_mem_usage=True
            )
            
            logger.info("✓ Mistral-7B loaded successfully")
            logger.info("🤖 Generating medical summary...")
            
            # Create medical report prompt with STRICT privacy rules
            prompt = f"""[INST] Medical report analyzer: Generate SHORT summary (1-2 paragraphs).

⚠️ NEVER INCLUDE:
❌ Names, age, gender
❌ Phone, address, location
❌ IDs, dates, report numbers

✅ ONLY INCLUDE:
• Test results (normal/abnormal)
• Health findings
• Recommendations

Medical Data:
{text[:2000]}

Generate summary with NO personal information.[/INST]"""
            
            inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=2048)
            
            # Move inputs to same device as model
            if torch.cuda.is_available():
                inputs = {k: v.cuda() for k, v in inputs.items()}
            
            logger.info("🤖 Generating summary with Mistral-7B...")
            with torch.no_grad():
                outputs = model.generate(
                    **inputs,
                    max_new_tokens=max_length,
                    temperature=0.7,
                    do_sample=True,
                    top_p=0.9,
                    repetition_penalty=1.2
                )
            
            summary = tokenizer.decode(outputs[0], skip_special_tokens=True)
            if "[/INST]" in summary:
                summary = summary.split("[/INST]")[-1].strip()
            
            logger.info("✓ Summary generated successfully")
            
            # STEP 2: Unload Mistral immediately after generation
            logger.info("🧹 Unloading Mistral-7B to free memory...")
            del model
            del tokenizer
            del inputs
            del outputs
            
            # Clear memory again
            self._clear_memory()
            logger.info("✓ Mistral-7B unloaded, memory freed")
            
            return summary
            
        except Exception as e:
            logger.error(f"✗ Mistral-7B error: {str(e)}")
            
            # Cleanup on error - be very thorough
            try:
                if model is not None:
                    del model
                if tokenizer is not None:
                    del tokenizer
                if 'inputs' in locals():
                    del inputs
                if 'outputs' in locals():
                    del outputs
                self._clear_memory()
            except Exception as cleanup_error:
                logger.warning(f"Cleanup warning: {str(cleanup_error)}")
            
            # Return None instead of error string to trigger fallback
            return None
    
    def _simple_summary(self, text: str, max_sentences: int = 5) -> str:
        """Simple extractive summarization fallback"""
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        
        summary_sentences = sentences[:max_sentences]
        summary = '. '.join(summary_sentences)
        
        if summary and not summary.endswith('.'):
            summary += '.'
            
        return summary
    
    # ============================================================================
    # Main execution method
    # ============================================================================
    def process_document(
        self, 
        file_path: str, 
        file_type: str,
        use_ai_summary: bool = True,
        progress_callback = None
    ) -> DocumentState:
        """
        Process document through the complete pipeline
        
        Args:
            file_path: Path to the document file
            file_type: Type of file (png, pdf, docx, txt)
            use_ai_summary: Whether to attempt AI summarization
            progress_callback: Optional callback for progress updates
            
        Returns:
            Final state with all processing results
        """
        
        # Initialize state
        initial_state = DocumentState(
            file_path=file_path,
            file_type=file_type,
            raw_text="",
            cleaned_text="",
            entities={},
            structured_data={},
            summary="",
            error="",
            processing_step="Starting...",
            use_ai_summary=use_ai_summary
        )
        
        # Execute the graph
        final_state = initial_state
        for step_state in self.graph.stream(initial_state):
            # Get the current node and state
            node_name = list(step_state.keys())[0]
            final_state = step_state[node_name]
            
            if progress_callback:
                progress_callback(node_name, final_state)
        
        # Return final state
        return final_state


# ============================================================================
# Utility: Model Status Checker (Isolated service)
# ============================================================================
class ModelStatusChecker:
    """Check availability of all models (microservice pattern)"""
    
    @staticmethod
    def check_all_models(models_dir: Path) -> dict:
        """Check which models are available"""
        status = {
            'tesseract': ModelStatusChecker._check_tesseract(),
            'spacy': ModelStatusChecker._check_spacy(),
            'regex': ModelStatusChecker._check_regex(models_dir),
            'mistral': ModelStatusChecker._check_mistral(models_dir)
        }
        return status
    
    @staticmethod
    def _check_tesseract() -> bool:
        try:
            import pytesseract
            paths = [
                r"D:\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
            ]
            return any(os.path.exists(p) for p in paths)
        except:
            return False
    
    @staticmethod
    def _check_spacy() -> bool:
        try:
            import spacy
            spacy.load("en_core_web_sm")
            return True
        except:
            return False
    
    @staticmethod
    def _check_regex(models_dir: Path) -> bool:
        return (models_dir / "regex_patterns.py").exists()
    
    @staticmethod
    def _check_mistral(models_dir: Path) -> bool:
        return (models_dir / "Mistral-7B-Instruct").exists()
